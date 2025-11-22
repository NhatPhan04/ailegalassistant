from __future__ import annotations

import os
import json
import logging
import pathlib
import re
from dataclasses import dataclass
from typing import List, Dict, Any, Optional, Tuple

# --- 3rd Party Libraries ---
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer, CrossEncoder
import faiss
import numpy as np
from rank_bm25 import BM25Okapi
from docx import Document
import google.generativeai as genai

# ===========================================================
# 0. CẤU HÌNH HỆ THỐNG & LOGGING
# ===========================================================

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(name)s: %(message)s",
)
logger = logging.getLogger("LegalAI")

BASE_DIR = pathlib.Path(__file__).parent
DATA_LAWS_DIR = BASE_DIR / "data_laws"
INDEX_DIR = BASE_DIR / "index_laws"
CONTRACT_DIR = BASE_DIR / "contracts"
# Cần tạo thêm 2 file này trong thư mục BASE_DIR
CHECKLIST_TEMPLATE_PATH = pathlib.Path(r"D:\Project\main\BE\check list\checklist_template.docx")
CHECKLIST_FINAL_PATH    = pathlib.Path(r"D:\Project\main\BE\check list\checklist_final.docx")


for d in [DATA_LAWS_DIR, INDEX_DIR, CONTRACT_DIR]:
    d.mkdir(exist_ok=True, parents=True)

load_dotenv()

GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
EMBED_MODEL_NAME = "sentence-transformers/all-MiniLM-L6-v2"
RERANK_MODEL_NAME = "cross-encoder/ms-marco-MiniLM-L-6-v2" # Model Re-ranking nhẹ
GCS_BUCKET_NAME = os.getenv("GCS_BUCKET_NAME")
GCS_LAWS_PREFIX = os.getenv("GCS_LAWS_PREFIX", "law/")

if not GEMINI_API_KEY:
    logger.warning("⚠️ CẢNH BÁO: GEMINI_API_KEY chưa được cấu hình.")

# ===========================================================
# 1. SYSTEM PROMPTS (CORE)
# ===========================================================

CORE_SYSTEM_PROMPT = """
# VAI TRÒ (ROLE)

Bạn là **AI Legal Assistant** – một mô hình ngôn ngữ lớn được huấn luyện chuyên sâu về **pháp lý doanh nghiệp tại Việt Nam**.

Bạn có ba nhiệm vụ chính:

1. TRA CỨU LUẬT THEO NGỮ NGHĨA (Semantic Legal Lookup)  
2. PHÂN TÍCH HỢP ĐỒNG DỰA TRÊN CHECKLIST DO HỆ THỐNG CUNG CẤP (Contract Analysis)  
3. LUẬT SƯ ONLINE – PHÂN TÍCH TÌNH HUỐNG THỰC TẾ CHO DOANH NGHIỆP (Virtual Corporate Lawyer)  

Bạn phải **trả lời bằng tiếng Việt**, rõ ràng, logic, có cấu trúc, **không bịa luật** và **chỉ dựa trên dữ liệu (CONTEXT) được cung cấp** hoặc phạm vi kiến thức pháp luật doanh nghiệp được cho phép.


────────────────────────────────────────
# PHẠM VI CHUYÊN MÔN (DOMAIN – IN SCOPE)

Bạn **được phép phân tích, tra cứu và giải thích pháp lý** trong phạm vi **pháp luật doanh nghiệp tại Việt Nam**, cụ thể:

## 1. Luật Đầu tư
- Luật Đầu tư 2020  
- Các văn bản sửa đổi, bổ sung  
- Nghị định hướng dẫn  
- Hình thức đầu tư, ưu đãi đầu tư, nghĩa vụ nhà đầu tư, ngành nghề cấm / hạn chế  

## 2. Luật Doanh nghiệp
- Thành lập doanh nghiệp  
- Loại hình doanh nghiệp (CTCP, TNHH, DNTN, công ty hợp danh…)  
- Quyền và nghĩa vụ của cổ đông, thành viên góp vốn  
- Cơ cấu quản trị, đại hội đồng cổ đông, hội đồng thành viên, HĐQT  
- Vốn điều lệ, góp vốn, chuyển nhượng vốn, điều lệ doanh nghiệp  

## 3. Luật Thương mại
- Hợp đồng mua bán hàng hóa B2B  
- Hợp đồng cung ứng dịch vụ  
- Đại lý, phân phối, nhượng quyền thương mại  
- Trung gian thương mại: môi giới, ủy thác, đại diện…  
- Đấu giá, đấu thầu, logistics  
- Vi phạm hợp đồng, phạt vi phạm, bồi thường thiệt hại  

## 4. Luật Quản lý Thuế
- Quản lý thuế đối với doanh nghiệp  
- Kê khai – nộp thuế – hoàn thuế  
- Xử lý vi phạm về thuế  
- Quản lý hóa đơn/chứng từ thuế  

## 5. Thuế Giá trị Gia tăng (VAT)
- Đối tượng chịu thuế – không chịu thuế  
- Thuế suất  
- Khấu trừ, hoàn thuế  
- Nghĩa vụ kê khai và thời điểm phát sinh nghĩa vụ  

## 6. Thuế Thu nhập Doanh nghiệp (TNDN)
- Cách xác định thu nhập chịu thuế  
- Chi phí được trừ – không được trừ  
- Thu nhập miễn thuế  
- Ưu đãi thuế TNDN, ưu đãi đầu tư  

## 7. Hóa đơn & Chứng từ
- Hóa đơn điện tử  
- Lập, điều chỉnh, thay thế, hủy hóa đơn  
- Trách nhiệm doanh nghiệp khi xuất hóa đơn  

## 8. Văn bản sửa đổi – bổ sung & Nghị định kèm theo
- Bạn phải áp dụng đúng nội dung văn bản trong CONTEXT hoặc tập luật được cung cấp.  
- Bạn **không được tự tạo “luật mới”, điều khoản mới** nếu không có trong dữ liệu hoặc không phù hợp với hệ thống pháp luật hiện hành.


────────────────────────────────────────
# PHẠM VI BỊ CẤM (OUT OF SCOPE)

Bạn **không được trả lời, tư vấn** các lĩnh vực sau, trừ khi câu hỏi gắn với pháp lý doanh nghiệp:

- Hình sự cá nhân  
- Dân sự cá nhân (vay mượn, tranh chấp nhỏ lẻ…)  
- Hôn nhân – gia đình  
- Đất đai, nhà ở, sổ đỏ, tài sản cá nhân  
- Thuế thu nhập cá nhân  
- Sức khỏe – y tế  
- Tín ngưỡng – tôn giáo  
- Tài chính cá nhân, chứng khoán cá nhân  

Nếu câu hỏi rơi hoàn toàn vào các lĩnh vực trên, bạn phải trả lời:

> “Câu hỏi này nằm ngoài phạm vi pháp lý doanh nghiệp mà tôi được phép hỗ trợ.”

Không được cố gắng tư vấn vượt phạm vi.


────────────────────────────────────────
# SỬ DỤNG CONTEXT & DỮ LIỆU (RAG BEHAVIOR)

Hệ thống có thể cung cấp cho bạn các phần dữ liệu sau:

- **CONTEXT_LUAT**: trích dẫn điều luật, nghị định, thông tư liên quan.  
- **CONTEXT_HOP_DONG**: nội dung hợp đồng hoặc trích đoạn hợp đồng.  
- **CONTEXT_CHECKLIST**: bộ checklist phân tích hợp đồng (20–50 mục).  
- **CONTEXT_TINH_HUONG**: mô tả tình huống kinh doanh thực tế của doanh nghiệp.  

Quy tắc:

1. **Ưu tiên tuyệt đối CONTEXT**  
   - Luôn đọc kỹ và sử dụng nội dung trong CONTEXT để trả lời.  
   - Nếu câu hỏi có thể trả lời bằng CONTEXT → chỉ dựa trên đó, giải thích lại cho rõ.

2. **Nếu CONTEXT không đủ hoặc không có**  
   - Bạn chỉ được sử dụng **kiến thức pháp luật doanh nghiệp Việt Nam trong phạm vi cho phép**.  
   - Nếu vẫn không đủ căn cứ để kết luận → bạn phải trả lời:  
     > “Dữ liệu không đủ để đưa ra kết luận chính xác.”  

3. **Tuyệt đối không bịa luật, không bịa số điều, khoản, điểm**  
   - Chỉ nhắc đến điều/khoản khi CONTEXT đã cung cấp hoặc đó là kiến thức pháp luật cơ bản, chắc chắn.  
   - Nếu không chắc, hãy nói rõ: “Không đủ dữ liệu để xác định chính xác điều/khoản cụ thể.”


────────────────────────────────────────
# 3 MODE HOẠT ĐỘNG CHÍNH

Hệ thống có thể truyền vào một tham số, ví dụ: `"mode": "tra_cuu" | "hop_dong" | "luat_su_online"`.

Bạn phải **điều chỉnh cách trả lời** tương ứng với mode.

────────────────────────
## MODE 1 – TRA CỨU LUẬT (SEMANTIC LEGAL LOOKUP)

Mục tiêu:
- Trả lời câu hỏi pháp lý dựa trên văn bản luật được cung cấp.  
- Giải thích luật bằng ngôn ngữ dễ hiểu cho doanh nghiệp.  
- Tóm tắt quy định, điều kiện, thủ tục.

Quy tắc:
- Dùng CONTEXT_LUAT để:
  - Tóm tắt nội dung chính.  
  - Giải thích ý nghĩa thực tiễn.  
- Không bịa số điều luật. Nếu trong CONTEXT không có số điều → không tự nghĩ thêm.

Cấu trúc trả lời (4 phần bắt buộc):
1) **Kết luận ngắn gọn**  
2) **Căn cứ pháp lý hoặc phân tích (theo CONTEXT_LUAT)**  
3) **Giải thích chi tiết / ví dụ thực tế (nếu phù hợp)**  
4) **Cảnh báo và gợi ý hành động cho doanh nghiệp**

────────────────────────
## MODE 2 – PHÂN TÍCH HỢP ĐỒNG (CONTRACT ANALYSIS)

Dữ liệu được cung cấp:
- **CONTEXT_HOP_DONG**: nội dung hợp đồng cần phân tích.  
- **CONTEXT_CHECKLIST**: danh sách các tiêu chí/điều khoản cần đối chiếu.  

Bạn **không được tự tạo checklist**; chỉ được phân tích dựa trên checklist do hệ thống cung cấp.

Mục tiêu:
- Tóm tắt hợp đồng và xác định loại hợp đồng.  
- Đối chiếu từng mục trong checklist với nội dung hợp đồng.  
- Phân loại từng vấn đề thành:
  1. Đã có và rõ ràng.  
  2. Có nhưng mơ hồ / không đầy đủ.  
  3. Thiếu hoàn toàn.  
  4. Có nhưng bất lợi cho doanh nghiệp người dùng.  
- Đánh giá mức độ rủi ro (Thấp – Trung bình – Cao).  
- Gợi ý chỉnh sửa, bổ sung điều khoản (dựa trên checklist).  
- Đề xuất câu hỏi nên hỏi lại đối tác.

Quy trình phân tích:
trước khi phan tích phải nhận dạng là hợp đồng final hay hợp đồng template 
**Bước 1 – Xác định loại hợp đồng**  
- Ví dụ: Hợp đồng dịch vụ, hợp đồng phân phối, mua bán B2B, đại lý, hợp tác kinh doanh...

**Bước 2 – Tóm tắt cấu trúc hợp đồng**  
- Tóm tắt ngắn gọn (3–7 dòng):  
  - Các bên trong hợp đồng  
  - Phạm vi công việc / đối tượng  
  - Thời hạn  
  - Giá & phương thức thanh toán  
  - Cam kết chính  
  - Điều khoản chấm dứt  
  - Trách nhiệm, bồi thường  

**Bước 3 – Đối chiếu checklist**  
Với mỗi mục trong CONTEXT_CHECKLIST:
- Kiểm tra xem hợp đồng có đề cập hay không.  
- Nếu có → đánh giá là:
  - Rõ ràng  
  - Mơ hồ / thiếu chi tiết  
  - Bất lợi cho doanh nghiệp  
- Nếu không có → đánh dấu “Thiếu”.  

Cách diễn đạt nên mang tính pháp lý chuyên nghiệp, ví dụ:
- “Điều khoản thanh toán không quy định thời hạn cụ thể, tiềm ẩn rủi ro chậm thanh toán.”  
- “Điều khoản chấm dứt không nêu rõ trường hợp đơn phương chấm dứt, rủi ro cho doanh nghiệp.”  
- “Đã có điều khoản bảo mật nhưng chưa quy định chế tài khi vi phạm.”

**Bước 4 – Đánh giá rủi ro**  
- Gán mức rủi ro cho từng nhóm điều khoản:  
  - Thấp: rõ ràng, cân bằng quyền lợi.  
  - Trung bình: có điểm mơ hồ hoặc thiếu sót nhưng có thể chấp nhận.  
  - Cao: thiếu điều khoản quan trọng hoặc bất lợi rõ rệt.

**Bước 5 – Đề xuất hành động**  
- Điều khoản nên bổ sung.  
- Điều khoản nên sửa hoặc đàm phán lại.  
- Câu hỏi nên gửi cho đối tác.  
- Cảnh báo rủi ro nếu giữ nguyên hợp đồng.

Cấu trúc câu trả lời (5 phần bắt buộc trong mode HỢP ĐỒNG):

1. **Tóm tắt hợp đồng (3–7 dòng)**  
2. **Đối chiếu checklist (liệt kê theo từng mục, phân nhóm ĐÃ CÓ / MƠ HỒ / THIẾU / BẤT LỢI)**  
3. **Đánh giá mức độ rủi ro (thấp – trung bình – cao)**  
4. **Gợi ý chỉnh sửa / bổ sung điều khoản**  
5. **Câu hỏi nên hỏi lại đối tác & cảnh báo quan trọng**

Lưu ý:
- Không được tự bịa checklist mới.  
- Không được khẳng định chắc chắn “an toàn 100%” hoặc “không có rủi ro”.  
- Chỉ phân tích trên tinh thần hỗ trợ doanh nghiệp hiểu rủi ro và chuẩn bị đàm phán.

────────────────────────
## MODE 3 – LUẬT SƯ ONLINE (VIRTUAL CORPORATE LAWYER)

Khi người dùng mô tả một **tình huống thực tế** (tranh chấp, rủi ro, vướng mắc trong hoạt động doanh nghiệp), bạn phải vận hành như một **luật sư doanh nghiệp tư vấn**.

Giọng điệu & phong cách:
- Chuyên nghiệp, điềm tĩnh, tự tin.  
- Trả lời thẳng vào trọng tâm.  
- Có quan điểm rõ ràng, có lập luận.  
- Cách nói như luật sư:  
  - “Dựa trên tình huống anh/chị cung cấp, tôi đánh giá như sau…”  
  - “Rủi ro lớn nhất trong vụ việc này nằm ở…”  
  - “Nếu ở vị trí doanh nghiệp, tôi sẽ đề xuất hướng xử lý như sau…”

Nhiệm vụ:
- Xác định vấn đề pháp lý cốt lõi trong tình huống.  
- Tóm tắt lại tình huống bằng ngôn ngữ pháp lý.  
- Phân tích rủi ro:  
  - Rủi ro pháp lý.  
  - Rủi ro thương mại.  
  - Rủi ro thực thi hợp đồng.  
  - Rủi ro về chứng cứ, hồ sơ.  
- Giải thích quyền và nghĩa vụ các bên.  
- Đưa ra **nhiều hướng xử lý** (thương lượng, thông báo, sửa hợp đồng, chuẩn bị tranh chấp…).  
- Đưa ra checklist hành động cụ thể cho doanh nghiệp.

Giới hạn:
- Không được tuyên bố “đảm bảo thắng kiện”, “100% đúng luật”.  
- Không đưa tư vấn mang tính cam kết pháp lý ràng buộc.  
- Luôn nhắc: nội dung chỉ là tham khảo, không thay thế ý kiến của luật sư hành nghề/thực tế.

Cấu trúc trả lời (5 phần bắt buộc):

1) **Đánh giá sơ bộ của luật sư**  
2) **Phân tích pháp lý và rủi ro**  
3) **Chiến lược xử lý (tối thiểu 2–4 hướng)**  
4) **Khuyến nghị thực tế (checklist hành động)**  
5) **Cảnh báo cần lưu ý và khuyến nghị tham khảo luật sư thực tế**


────────────────────────────────────────
# STYLE & BEHAVIOR – CÁCH TRẢ LỜI CHUNG

Dù ở mode nào, bạn vẫn phải tuân theo **khung 4 phần** dưới đây (có thể lồng với cấu trúc riêng từng mode):

1) **Kết luận ngắn gọn**  
   - Một – hai câu nêu kết luận chính hoặc đánh giá sơ bộ.  

2) **Căn cứ pháp lý hoặc phân tích (theo CONTEXT)**  
   - Dẫn chiếu hoặc tóm tắt điều luật/điều khoản/ý chính của CONTEXT.  

3) **Giải thích chi tiết / phân tích từng điểm**  
   - Trình bày rõ ràng, có lý do, có ví dụ nếu cần.  

4) **Cảnh báo và gợi ý hành động**  
   - Chỉ ra rủi ro, điểm cần cẩn trọng.  
   - Đề xuất các bước tiếp theo hoặc câu hỏi nên chuẩn bị.  

Nguyên tắc chung:
- Không dùng ngôn ngữ mơ hồ, vòng vo.  
- Không dùng các khẳng định tuyệt đối (“chắc chắn 100%”, “không có rủi ro…”).  
- Không được bịa thông tin pháp lý.  
- Nếu thiếu dữ liệu → nói rõ:  
  > “Dữ liệu không đủ để đưa ra kết luận chính xác.”  

────────────────────────────────────────
# NĂNG LỰC MONG ĐỢI SAU HUẤN LUYỆN

Mô hình sau khi huấn luyện phải có khả năng:

- Hiểu đúng ngữ nghĩa câu hỏi pháp lý doanh nghiệp.  
- Phân loại đúng mode: Tra cứu / Hợp đồng / Luật sư online (khi hệ thống cung cấp mode).  
- Đọc và trích xuất thông tin quan trọng từ hợp đồng.  
- Đối chiếu hợp đồng với checklist một cách có hệ thống.  
- Đánh giá và “chấm điểm” mức độ đầy đủ / an toàn (nếu được yêu cầu).  
- Phát hiện và nêu rõ rủi ro pháp lý, thương mại trong từng tình huống.  
- Giải thích luật bằng ngôn ngữ dễ hiểu cho doanh nghiệp.  
- Đưa ra gợi ý hành động thực tế, nhưng luôn giữ giới hạn: chỉ mang tính tham khảo, không thay thế luật sư hành nghề.

────────────────────────────────────────
# FINAL RULE

Khi CONTEXT hoặc thông tin tình huống **không đủ rõ** để kết luận:

- Bạn **phải từ chối kết luận dứt khoát** và nói:  
  > “Dữ liệu không đủ để đưa ra kết luận chính xác.”  

- Bạn có thể:
  - Gợi ý thêm loại thông tin/hồ sơ mà doanh nghiệp cần cung cấp.  
  - Đưa ra các hướng suy nghĩ, nhưng không khẳng định đây là kết luận cuối cùng.  

Tuyệt đối **không được bịa luật, không được suy đoán quá phạm vi dữ liệu** và **không được vượt ra ngoài phạm vi pháp lý doanh nghiệp.**

"""

# ===========================================================
# 2. UTILS & GEMINI CLIENT
# ===========================================================

class GeminiClient:
    _model: Optional[genai.GenerativeModel] = None

    @classmethod
    def get_model(cls) -> genai.GenerativeModel:
        if cls._model is None:
            if not GEMINI_API_KEY:
                raise RuntimeError("Thiếu GEMINI_API_KEY")
            genai.configure(api_key=GEMINI_API_KEY)
            cls._model = genai.GenerativeModel("gemini-2.5-flash")
        return cls._model

    @classmethod
    def generate_text(cls, prompt: str) -> str:
        try:
            return cls.get_model().generate_content(prompt).text.strip()
        except Exception as e:
            logger.error(f"Gemini Error: {e}")
            return ""

    @classmethod
    def generate_json(cls, prompt: str, fallback: Any) -> Any:
        try:
            resp = cls.get_model().generate_content(
                prompt,
                generation_config=genai.GenerationConfig(response_mime_type="application/json")
            )
            return json.loads(resp.text)
        except Exception as e:
            logger.error(f"JSON Error: {e}")
            return fallback


def detect_contract_status(text: str) -> Dict:
    """
    Phân loại TEMPLATE hay FINAL dựa trên rule-based trước,
    LLM chỉ dùng khi không chắc chắn.
    """

    # --- RULE-BASED TEMPLATE CHECK ---
    template_patterns = [
        r"\.{3,}",             # ......
        r"_ {3,}|_{3,}",       # ___
        r"\[.*?\]",            # [Tên Bên A]
        r"{.*?}",              # {Ngày}
        r"<.*?>",              # <Placeholder>
        r"Điền vào",           # hướng dẫn điền mẫu
        r"\( *\) *Có",         # ( ) Có
        r"\( *\) *Không",
        r"…+",                 # dấu ba chấm unicode
    ]

    template_hits = sum(1 for p in template_patterns if re.search(p, text, flags=re.IGNORECASE))

    # --- RULE-BASED FINAL CHECK ---
    final_patterns = [
        r"\b\d{1,2}/\d{1,2}/\d{2,4}\b",                  # 12/01/2024
        r"ngày\s+\d{1,2}\s+tháng\s+\d{1,2}\s+năm\s+\d{4}",
        r"Mã số thuế|MST",                               # MST
        r"ông\s+[A-ZÀÁẠẢÃÈÉẺẸÊẾỀỂỆÔỒỐỔỘƯỨỪỰỬ]",         # tên người đại diện
        r"Công ty TNHH|Công ty Cổ phần|CTCP|TNHH",
        r"\d{1,3}(?:\.\d{3})+",                          # số tiền: 1.500.000
    ]

    final_hits = sum(1 for p in final_patterns if re.search(p, text, flags=re.IGNORECASE))

    # --- RULE-BASED DECISION ---
    if template_hits >= 3 and final_hits < 2:
        return {"status": "TEMPLATE", "reason": "Phát hiện nhiều placeholder, chưa điền dữ liệu."}

    if final_hits >= 3 and template_hits <= 1:
        return {"status": "FINAL", "reason": "Thông tin đã điền đầy đủ: ngày, MST, doanh nghiệp, số tiền."}

    # --- FALLBACK TO LLM ---
    prompt = f"""
    Bạn là chuyên gia phân loại hợp đồng.
    Nếu văn bản có nhiều dấu "....", "__", "[]", "<>" → TEMPLATE.
    Nếu thông tin tên công ty, MST, số tiền, ngày tháng được điền đầy đủ → FINAL.

    Văn bản:
    {text[:5000]}

    Hãy trả về JSON:
    {{
        "status": "TEMPLATE" hoặc "FINAL",
        "reason": "Giải thích ngắn"
    }}
    """

    result = GeminiClient.generate_json(prompt, fallback={
        "status": "UNKNOWN",
        "reason": "Không phân loại được"
    })
    return result


def read_docx(path: pathlib.Path) -> str:
    """Đọc DOCX (Text + Table)"""
    try:
        if not path.exists():
            return ""
        doc = Document(str(path))
        full_text = []
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    full_text.append(" | ".join(cells))
        return "\n".join(full_text)
    except Exception as e:
        logger.error(f"read_docx error: {e}")
        return ""


def chunk_law_text(text: str, source_name: str, min_len=20, max_chunk_size=4500) -> List[str]:
    """
    Cắt luật theo Điều + Metadata Injection 
    Thêm: [NGUỒN: Tên_File] vào đầu mỗi chunk.
    """
    article_pattern = r'(?:\n|^)(Điều\s+\d+[.:])'
    matches = list(re.finditer(article_pattern, text))
    chunks = []

    def inject_metadata(content: str) -> str:
        return f"[NGUỒN: {source_name}]\n{content}"

    if not matches:
        raw_chunks = [p.strip() for p in text.split("\n\n") if len(p.strip()) > min_len]
        return [inject_metadata(c) for c in raw_chunks]

    for i, match in enumerate(matches):
        start = match.start()
        end = matches[i+1].start() if i+1 < len(matches) else len(text)
        full_article = text[start:end].strip()

        if len(full_article) > max_chunk_size:
            header_match = re.match(r'(Điều\s+\d+[.:].*?)(\n|$)', full_article)
            header = header_match.group(1) if header_match else ""
            sub_parts = full_article.split("\n")
            current = ""
            for part in sub_parts:
                if len(current) + len(part) > max_chunk_size:
                    if current:
                        chunks.append(inject_metadata(current))
                    current = f"{header} (tiếp)... \n{part}"
                else:
                    current += f"\n{part}"
            if current:
                chunks.append(inject_metadata(current))
        else:
            if len(full_article) >= min_len:
                chunks.append(inject_metadata(full_article))

    return chunks


def download_law_docs_from_gcs():
    """Tải file từ GCS với log chi tiết (nếu có cấu hình)."""
    bucket_name = os.getenv("GCS_BUCKET_NAME")
    prefix = os.getenv("GCS_LAWS_PREFIX")

    if not bucket_name:
        logger.info("ℹ️ GCS_BUCKET_NAME chưa cấu hình -> Bỏ qua tải từ Cloud.")
        return

    try:
        from google.cloud import storage
        from google.oauth2 import service_account

        KEY_PATH = BASE_DIR / "gcs_key.json"

        if KEY_PATH.exists():
            logger.info(f"🔑 Tìm thấy key GCS tại: {KEY_PATH}")
            credentials = service_account.Credentials.from_service_account_file(str(KEY_PATH))
            client = storage.Client(credentials=credentials)
        else:
            logger.warning("⚠️ Không thấy file gcs_key.json. Thử dùng credentials mặc định...")
            client = storage.Client()

        bucket = client.bucket(bucket_name)
        blobs = bucket.list_blobs(prefix=prefix)

        count = 0
        logger.info(f"📡 Đang kết nối GCS Bucket: {bucket_name}...")

        for blob in blobs:
            if blob.name.lower().endswith(".docx"):
                filename = blob.name.split("/")[-1]
                if not filename:
                    continue
                local_path = DATA_LAWS_DIR / filename
                if not local_path.exists() or local_path.stat().st_size == 0:
                    logger.info(f"⬇️ Đang tải: {filename}")
                    blob.download_to_filename(str(local_path))
                    count += 1

        if count > 0:
            logger.info(f"✅ Đã tải {count} file mới từ Cloud.")
        else:
            logger.info("⚡ Dữ liệu local đã đồng bộ.")

    except Exception as e:
        logger.error(f"❌ GCS Error: {e}")
        logger.error("👉 Gợi ý: Kiểm tra Service Account hoặc cấu hình GCS.")


# ===========================================================
# 3. ADVANCED VECTOR STORE (HYBRID + VALIDITY FILTER)
# ===========================================================

@dataclass
class LawChunk:
    text: str
    source_file: str


class LawVectorStore:
    """
    Store tích hợp:
    1. Validity Filter: Loại bỏ luật năm cũ.
    2. Hybrid Search: Vector (FAISS) + Keyword (BM25).
    3. Re-ranking: Cross-Encoder.
    """
    def __init__(self):
        self.embedder = SentenceTransformer(EMBED_MODEL_NAME)
        self.cross_encoder = CrossEncoder(RERANK_MODEL_NAME)
        self.index = None
        self.chunks: List[LawChunk] = []
        self.bm25 = None  # Keyword search engine

    def _filter_valid_laws(self, dir_path: pathlib.Path) -> List[pathlib.Path]:
        files = list(dir_path.glob("*.docx"))
        law_map = {}
        pattern = r"(.+?)_(\d{4})"

        valid_files = []
        for f in files:
            match = re.search(pattern, f.stem)
            if match:
                name_core = match.group(1)
                year = int(match.group(2))

                if name_core not in law_map:
                    law_map[name_core] = (year, f)
                else:
                    if year > law_map[name_core][0]:
                        law_map[name_core] = (year, f)
            else:
                valid_files.append(f)

        for _, val in law_map.items():
            valid_files.append(val[1])

        logger.info(f"🧹 Lọc luật cũ: {len(files)} -> {len(valid_files)} file hiệu lực.")
        return valid_files

    def build(self):
        valid_files = self._filter_valid_laws(DATA_LAWS_DIR)
        if not valid_files:
            logger.warning("⚠️ Không có file dữ liệu.")
            return

        all_chunks = []
        for f in valid_files:
            text = read_docx(f)
            chunks = chunk_law_text(text, f.name)
            for c in chunks:
                all_chunks.append(LawChunk(text=c, source_file=f.name))

        if not all_chunks:
            return

        self.chunks = all_chunks

        # Build FAISS
        logger.info("⚡ Building FAISS Index...")
        embeddings = self.embedder.encode([c.text for c in all_chunks], convert_to_numpy=True)
        self.index = faiss.IndexFlatIP(embeddings.shape[1])
        self.index.add(embeddings)

        # Build BM25
        logger.info("🔑 Building BM25 Index...")
        tokenized_corpus = [c.text.lower().split() for c in all_chunks]
        self.bm25 = BM25Okapi(tokenized_corpus)

        logger.info(f"✅ Index xong {len(all_chunks)} chunks.")
        self.save()

    def hybrid_search(self, query: str, top_k=50, final_k=5) -> List[LawChunk]:
        if not self.chunks or self.index is None or self.bm25 is None:
            return []

        # Semantic search
        q_vec = self.embedder.encode([query], convert_to_numpy=True)
        _, v_idxs = self.index.search(q_vec, top_k)
        vector_results = {idx for idx in v_idxs[0] if 0 <= idx < len(self.chunks)}

        # BM25
        tokenized_query = query.lower().split()
        bm25_top = self.bm25.get_top_n(tokenized_query, self.chunks, n=top_k)

        candidate_chunks: List[LawChunk] = []
        for idx in vector_results:
            candidate_chunks.append(self.chunks[idx])
        for c in bm25_top:
            if c not in candidate_chunks:
                candidate_chunks.append(c)

        if not candidate_chunks:
            return []

        pairs = [[query, c.text] for c in candidate_chunks]
        scores = self.cross_encoder.predict(pairs)
        sorted_indices = np.argsort(scores)[::-1]

        final_results = []
        for i in range(min(final_k, len(candidate_chunks))):
            idx = sorted_indices[i]
            final_results.append(candidate_chunks[idx])

        return final_results

    def save(self):
        if self.index is None:
            return

        INDEX_DIR.mkdir(exist_ok=True, parents=True)

        faiss.write_index(self.index, str(INDEX_DIR / "laws.faiss"))

        with (INDEX_DIR / "laws_meta.jsonl").open("w", encoding="utf-8") as f:
            for c in self.chunks:
                data = {"text": c.text, "source_file": c.source_file}
                f.write(json.dumps(data, ensure_ascii=False) + "\n")

        logger.info("💾 Đã lưu Index xuống ổ cứng.")

    def load(self) -> bool:
        if not (INDEX_DIR / "laws.faiss").exists():
            return False

        logger.info("📂 Đang load Index từ ổ cứng...")

        self.index = faiss.read_index(str(INDEX_DIR / "laws.faiss"))

        self.chunks = []
        with (INDEX_DIR / "laws_meta.jsonl").open("r", encoding="utf-8") as f:
            for line in f:
                data = json.loads(line)
                self.chunks.append(LawChunk(text=data["text"], source_file=data["source_file"]))

        if self.chunks:
            tokenized = [c.text.lower().split() for c in self.chunks]
            self.bm25 = BM25Okapi(tokenized)

        logger.info(f"✅ Đã load {len(self.chunks)} chunks.")
        return True


# ===========================================================
# 4. AGENTS (UPGRADED LOGIC)
# ===========================================================

class IntentNormalizationAgent:
    """Agent 1: Hybrid (LLM + Keyword Force)"""
    PROMPT = """
    Phân loại ý định user vào: 
    - "tra_cuu_luat": hỏi thủ tục, luật, hồ sơ.
    - "phan_tich_hop_dong": nhờ check file.
    - "goi_y_dieu_khoan": nhờ soạn thảo.
    - "chatchit": chào hỏi xã giao.
    Input: "{input}"
    Output JSON: {{ "clean_text": "...", "mode": "..." }}
    """

    def run(self, text: str) -> Dict:
        res = GeminiClient.generate_json(
            self.PROMPT.format(input=text),
            fallback={"clean_text": text, "mode": "tra_cuu_luat"}
        )
        keys = ["thủ tục", "đăng ký", "luật", "hồ sơ", "thuế", "cần gì", "như thế nào"]
        if res.get("mode") == "chatchit" and any(k in res.get("clean_text", "").lower() for k in keys):
            res["mode"] = "tra_cuu_luat"
        return res


class RAGRetrievalAgent:
    """
    Agent 2: CoT Framework 
    Tách câu hỏi -> Search 3 lần -> Gộp kết quả.
    """
    def __init__(self, store: LawVectorStore):
        self.store = store

    def run(self, complex_query: str) -> List[LawChunk]:
        prompt = f"""
        Phân tích câu hỏi: "{complex_query}"
        Hãy tách thành 3 search queries ngắn gọn để tìm kiếm trong luật.
        Output JSON list: ["query1", "query2", "query3"]
        """
        queries = GeminiClient.generate_json(prompt, fallback=[complex_query])
        if not isinstance(queries, list):
            queries = [complex_query]

        logger.info(f"🧠 CoT Queries: {queries}")

        all_results: List[LawChunk] = []
        for q in queries:
            results = self.store.hybrid_search(q, top_k=30, final_k=3)
            all_results.extend(results)

        seen = set()
        unique_results: List[LawChunk] = []
        for r in all_results:
            if r.text not in seen:
                unique_results.append(r)
                seen.add(r.text)

        return unique_results


class ContractAnalyzerAgent:
    """
    Agent phân tích hợp đồng:
    - Tự phân loại TEMPLATE / FINAL
    - Chọn checklist tương ứng
    """

    def __init__(self):
        self.checklist_template = read_docx(CHECKLIST_TEMPLATE_PATH)
        if not self.checklist_template:
            self.checklist_template = (
                "TIÊU CHUẨN HỢP ĐỒNG MẪU (DEFAULT):\n"
                "1. Các chỗ trống (placeholder) cần rõ ràng.\n"
                "2. Không có điều khoản trái luật.\n"
                "3. Hướng dẫn điền thông tin đầy đủ."
            )
            logger.warning(f"⚠️ Không đọc được {CHECKLIST_TEMPLATE_PATH}, dùng checklist mặc định.")

        self.checklist_final = read_docx(CHECKLIST_FINAL_PATH)
        if not self.checklist_final:
            self.checklist_final = (
                "TIÊU CHUẨN HỢP ĐỒNG FINAL (DEFAULT):\n"
                "1. Thông tin các bên đầy đủ (MST, Địa chỉ...).\n"
                "2. Điều khoản thanh toán, phạt vi phạm rõ ràng.\n"
                "3. Quyền và nghĩa vụ cân bằng."
            )
            logger.warning(f"⚠️ Không đọc được {CHECKLIST_FINAL_PATH}, dùng checklist mặc định.")

    def analyze(self, contract_text: str, store: Optional[LawVectorStore] = None) -> str:
        if not contract_text:
            return "❌ Lỗi: Không đọc được nội dung hợp đồng."

        status_info = detect_contract_status(contract_text)
        doc_type = status_info.get("status", "FINAL")
        reason = status_info.get("reason", "")
        logger.info(f"[ContractAnalyzer] Phát hiện loại hợp đồng: {doc_type} | Lý do: {reason}")

        if doc_type == "TEMPLATE":
            selected_checklist = self.checklist_template
            system_instruction = """
            ⚠️ PHÁT HIỆN: HỢP ĐỒNG MẪU (TEMPLATE).
            NHIỆM VỤ:
            1. Kiểm tra chất lượng mẫu có đúng chuẩn pháp lý không.
            2. Liệt kê tất cả các placeholder cần điền + rủi ro nếu điền sai.
            3. Cảnh báo các điều khoản còn thiếu so với checklist mẫu.
            4. Đề xuất bổ sung điều khoản quan trọng cho mẫu.
            """
        else:
            selected_checklist = self.checklist_final
            system_instruction = """
            ✅ PHÁT HIỆN: HỢP ĐỒNG ĐÃ ĐIỀN ĐẦY ĐỦ (FINAL/EXECUTED).
            NHIỆM VỤ:
            1. Xác định rủi ro pháp lý thực tế cho từng bên dựa trên thông tin đã điền.
            2. Kiểm tra tính hợp lệ của thông tin, số liệu, thời hạn, phạt.
            3. Đối chiếu checklist hoàn thiện để tìm mục thiếu / bất lợi.
            4. Đưa ra các vấn đề trọng yếu cần đàm phán lại.
            """

        if store:
            query = contract_text[:1500].replace("\n", " ")
            law_chunks = store.hybrid_search(query, top_k=40, final_k=8)
            law_block = "\n".join([f"- [Nguồn: {c.source_file}] {c.text[:500]}" for c in law_chunks])
        else:
            law_block = "Không sử dụng RAG."

        prompt = f"""
        {CORE_SYSTEM_PROMPT}

        {system_instruction}

        === PHÂN LOẠI ĐẦU VÀO ===
        - Loại văn bản: {doc_type}
        - Nhận định hệ thống: {reason}

        === DỮ LIỆU HỖ TRỢ ===
        • CHECKLIST ÁP DỤNG:  
        {selected_checklist}

        • LUẬT THAM CHIẾU (RAG):  
        {law_block}

        • NỘI DUNG HỢP ĐỒNG CẦN CHECK:  
        {contract_text[:30000]}

        =====================================================
        🎯 YÊU CẦU OUTPUT (THEO ĐÚNG CẤU TRÚC MARKDOWN)
        =====================================================
        Lưu ý: Loại bỏ ký tự đặc biệt, xuống dòng thừa. Format chuyên nghiệp.

        # 1. NHẬN DIỆN TÀI LIỆU  
        - Loại hợp đồng: {doc_type}  
        - Tóm tắt nội dung chính (3–7 dòng)

        # 2. ĐỐI CHIẾU CHECKLIST (Bảng chi tiết)

        | Mục Checklist | Đã có | Mơ hồ | Thiếu | Bất lợi | Ghi chú |
        |---------------|-------|--------|--------|---------|---------|

        # 3. PHÂN TÍCH RỦI RO (Tham chiếu điều luật rõ ràng)
        Với mỗi rủi ro:
        - Mô tả vấn đề
        - Điều khoản gây rủi ro trong hợp đồng
        - Căn cứ pháp lý (nếu có trong RAG)
        - Mức độ nghiêm trọng (Thấp / TB / Cao)
        - Tác động cụ thể lên doanh nghiệp

        # 4. GỢI Ý TỐI ƯU (Điều khoản nên sửa và lý do)
        - Liệt kê điểm cần sửa  
        - Đề xuất câu chữ mẫu (Drafting)
        - Gợi ý câu hỏi nên hỏi đối tác  

        # 5. CHẤM ĐIỂM HỢP ĐỒNG (0–100)

        Hãy chấm điểm theo bảng bên dưới và **thay thế toàn bộ `<...>` bằng giá trị thực** (KHÔNG để dấu `<` `>` trong output):

        ### 5. Điểm số hợp đồng

        | Tiêu chí                        | Điểm (0–10)      | Ghi chú ngắn gọn                        |
        |---------------------------------|------------------|----------------------------------------|
        | Độ rõ ràng (Clarity)            | <clarity>        | Ví dụ: Điều khoản rõ / còn mơ hồ       |
        | Cân bằng lợi ích (Balance)      | <balance>        | Ví dụ: Thiên lệch cho bên nào không    |
        | Rủi ro pháp lý (Risk)           | <risk>           | Điểm cao = rủi ro nhiều                |
        | **Điểm tổng hợp (Contract Score)** | **<contract_score>** | Trung bình sau khi xem xét các tiêu chí |

        **Mức độ rủi ro tổng thể:** **<THẤP / TRUNG BÌNH / CAO>**
        Lời khuyên: (Nếu dưới 70 điểm, yêu cầu người dùng xem xét kỹ lưỡng và chỉnh sửa lại hợp đồng trước khi ký kết).
        """

        return GeminiClient.generate_text(prompt)

    def suggest(self, req: str) -> str:
        return GeminiClient.generate_text(f"Soạn điều khoản phù hợp cho hợp đồng doanh nghiệp: {req}")


class LegalAnswerAgent:
    """
    Agent sinh câu trả lời cuối cùng cho:
    - tra_cuu_luat
    - luat_su_online
    (và có thể mở rộng sau)
    """

    def run(self, query: str, context: str, mode: str) -> str:
        if mode == "tra_cuu_luat":
            mode_instruction = """
            Bạn đang ở MODE: TRA CỨU LUẬT (SEMANTIC LEGAL LOOKUP).

            NHIỆM VỤ:
            - Dùng CONTEXT_LUAT bên dưới như nguồn chính để trả lời.
            - Nếu context có nội dung:
                + Tóm tắt quy định chính.
                + Giải thích ý nghĩa cho doanh nghiệp.
            - Nếu context là 'KHÔNG TÌM THẤY DỮ LIỆU...' thì:
                + Trả lời dựa trên kiến thức pháp luật doanh nghiệp chung.
                + Nhưng phải nói rõ: dữ liệu không đầy đủ, chỉ mang tính tham khảo.

            CẤU TRÚC TRẢ LỜI BẮT BUỘC:
            1) Kết luận ngắn gọn.
            2) Căn cứ pháp lý hoặc phân tích theo CONTEXT_LUAT.
            3) Giải thích chi tiết / ví dụ thực tế (nếu phù hợp).
            4) Cảnh báo và gợi ý hành động cho doanh nghiệp.
            """
        elif mode == "luat_su_online":
            mode_instruction = """
            Bạn đang ở MODE: LUẬT SƯ ONLINE (VIRTUAL CORPORATE LAWYER).

            NHIỆM VỤ:
            - Xem câu hỏi của user như một tình huống thực tế doanh nghiệp.
            - Dùng CONTEXT_LUAT (nếu có) để tham chiếu.
            - Phân tích rủi ro pháp lý + thương mại một cách thực tế.

            CẤU TRÚC TRẢ LỜI BẮT BUỘC:
            1) Đánh giá sơ bộ của luật sư.
            2) Phân tích pháp lý và rủi ro (tham chiếu context nếu có).
            3) Chiến lược xử lý (2–4 hướng).
            4) Checklist hành động cụ thể cho doanh nghiệp.
            5) Cảnh báo cần lưu ý và khuyến nghị tham khảo luật sư thực tế.
            """
        else:
            mode_instruction = """
            MODE không xác định rõ (fallback).
            Hãy trả lời theo phong cách AI Legal Assistant,
            giữ đúng khung 4 phần:
            1) Kết luận ngắn gọn
            2) Phân tích / Căn cứ
            3) Giải thích chi tiết
            4) Cảnh báo và gợi ý hành động
            """

        prompt = f"""
{CORE_SYSTEM_PROMPT}

================= NGỮ CẢNH (CONTEXT_LUAT / RAG) =================
{context}

================= CÂU HỎI CỦA NGƯỜI DÙNG =================
{query}

================= HƯỚNG DẪN MODE =================
{mode_instruction}

LƯU Ý:
- Không được bịa luật, không bịa điều/khoản nếu không có trong context hoặc kiến thức chắc chắn.
- Nếu context trống hoặc yếu, phải nói rõ: "Dữ liệu không đủ để đưa ra kết luận chính xác."
- Luôn trả lời bằng tiếng Việt, rõ ràng, có cấu trúc.
"""
        return GeminiClient.generate_text(prompt)


# ===========================================================
# 5. ORCHESTRATOR
# ===========================================================

class LegalOrchestrator:
    def __init__(self):
        logger.info("🚀 System Init...")
        download_law_docs_from_gcs()
        self.store = LawVectorStore()
        self.store.build()

        self.intent_agent = IntentNormalizationAgent()
        self.rag_agent = RAGRetrievalAgent(self.store)
        self.contract_agent = ContractAnalyzerAgent()
        self.answer_agent = LegalAnswerAgent()

    def process(self, user_input: str, file_path: str = None) -> str:
        try:
            intent = self.intent_agent.run(user_input)
            mode = intent["mode"]
            query = intent["clean_text"]

            logger.info(f"🔍 Process | Mode: {mode} | Query: {query}")

            # A: TRA CỨU LUẬT / LUẬT SƯ ONLINE
            if mode in ["tra_cuu_luat", "luat_su_online"]:
                chunks = self.rag_agent.run(query)

                print(f"\n[DEBUG] RAG tìm thấy: {len(chunks)} đoạn văn bản.")
                for i, c in enumerate(chunks[:3]):
                    print(f"  -> [{c.source_file}] {c.text[:50]}...")

                if chunks:
                    ctx = "\n\n".join([c.text for c in chunks])
                else:
                    logger.warning("⚠️ RAG trả về rỗng. AI sẽ trả lời dựa trên kiến thức nền kèm cảnh báo.")
                    ctx = "KHÔNG TÌM THẤY DỮ LIỆU TRONG CƠ SỞ DỮ LIỆU NỘI BỘ."

                return self.answer_agent.run(query, ctx, mode)

            # B: PHÂN TÍCH HỢP ĐỒNG
            elif mode == "phan_tich_hop_dong":
                if not file_path:
                    return (
                        "⚠️ **Thiếu file hợp đồng!**\n"
                        "Để tôi phân tích, bạn vui lòng nhập lại theo cú pháp:\n"
                        "> `file: đường/dẫn/đến/hop_dong.docx`"
                    )

                path_obj = pathlib.Path(file_path)
                if not path_obj.exists():
                    return f"❌ Lỗi: Không tìm thấy file tại đường dẫn: `{file_path}`"

                contract_text = read_docx(path_obj)
                if not contract_text:
                    return "❌ Lỗi: File rỗng hoặc không đọc được nội dung."

                logger.info(f"📄 Đang phân tích hợp đồng: {path_obj.name}")
                # Nếu muốn dùng RAG cho phân tích hợp đồng: truyền self.store
                return self.contract_agent.analyze(contract_text, store=self.store)

            # C: GỢI Ý / SOẠN THẢO ĐIỀU KHOẢN
            elif mode == "goi_y_dieu_khoan":
                logger.info("✍️ Đang soạn thảo điều khoản...")
                return self.contract_agent.suggest(query)

            # D: CHATCHIT (XÃ GIAO)
            elif mode == "chatchit":
                chat_prompt = f"""
                {CORE_SYSTEM_PROMPT}
                
                BỐI CẢNH: Người dùng đang giao tiếp xã giao (Chào hỏi/Hỏi danh tính).
                CÂU NÓI CỦA USER: "{query}"
                
                NHIỆM VỤ:
                1. Trả lời trực tiếp, thân thiện, ngắn gọn.
                2. KHÔNG đưa ra lời khuyên kỹ năng mềm (Ví dụ: KHÔNG nói "Bạn có thể trả lời là...").
                3. Luôn giữ vai là **AI Legal Assistant** chuyên về Pháp lý Doanh nghiệp.
                4. Nếu user hỏi "Bạn là ai?", hãy giới thiệu ngắn gọn về khả năng: Tra cứu luật, Soát xét hợp đồng, Tư vấn rủi ro.
                """
                return GeminiClient.generate_text(chat_prompt)

            # E: FALLBACK
            return (
                "Xin lỗi, tôi chưa hiểu rõ yêu cầu của bạn.\n"
                "Bạn có thể hỏi lại cụ thể hơn, ví dụ:\n"
                "- 'Thủ tục thành lập công ty TNHH?'\n"
                "- 'Soạn giúp tôi điều khoản bảo mật thông tin.'"
            )

        except Exception as e:
            logger.error(f"CRITICAL ERROR in Process: {e}")
            return f"⚠️ Hệ thống gặp lỗi kỹ thuật không mong muốn: {str(e)}"


if __name__ == "__main__":
    app = LegalOrchestrator()
    print("\n✅ System Ready (Advanced Mode)") 
    while True:
        try:
            u = input("\nYou: ").strip()
            if u in ["exit", "quit"]:
                break
            if not u:
                continue
            f = None
            if u.startswith("file:"):
                f = u.split(":", 1)[1].strip().replace('"', "")
                u = "Phân tích file"
            print(f"AI: {app.process(u, f)}")
        except Exception as e:
            print(f"Error: {e}")
